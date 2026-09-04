"""Generate the PC6M-10 software design and development document.

Usage:
    python tools/docs/generate_pc6m_software_design_doc.py
    python tools/docs/generate_pc6m_software_design_doc.py --output path/to/output.docx

Requires python-docx. The bundled Codex document runtime may be used when the
system Python environment does not provide that package.
"""

import argparse

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUT = REPO_ROOT / "docs" / "project" / "PC6M10_软件设计及开发文档_V1.1.docx"
parser = argparse.ArgumentParser(description="生成 PC6M-10 软件设计及开发文档")
parser.add_argument("--output", type=Path, default=DEFAULT_OUT, help="DOCX 输出路径")
OUT = parser.parse_args().output.resolve()
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
WHITE = "FFFFFF"
BLACK = "111827"
GOLD = "9A6B16"
RED = "9B1C1C"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)
sec.different_first_page_header_footer = True


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_explicit_fonts(rpr, east, latin=None):
    """Set OOXML fonts explicitly and remove theme fallbacks such as MS Gothic."""
    latin = latin or east
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in rfonts.attrib:
            del rfonts.attrib[key]
    for attr, value in (
        ("ascii", latin), ("hAnsi", latin), ("eastAsia", east), ("cs", latin)
    ):
        rfonts.set(qn(f"w:{attr}"), value)


def font_run(run, size=12, bold=False, color=BLACK, italic=False, east="宋体", latin="SimSun"):
    run.font.name = latin
    set_explicit_fonts(run._element.get_or_add_rPr(), east, latin)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


styles = doc.styles
normal = styles["Normal"]
normal.font.name = "SimSun"
set_explicit_fonts(normal._element.get_or_add_rPr(), "宋体", "SimSun")
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.first_line_indent = Pt(24)

for name, size, color, before, after in [
    ("Heading 1", 15, BLACK, 18, 9),
    ("Heading 2", 14, BLACK, 9, 9),
    ("Heading 3", 12, BLACK, 5.5, 0),
]:
    st = styles[name]
    st.font.name = "黑体"
    set_explicit_fonts(st._element.get_or_add_rPr(), "黑体", "黑体")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.first_line_indent = Pt(0)
    st.paragraph_format.keep_with_next = True

if "Heading 4" in styles:
    st = styles["Heading 4"]
    st.font.name = "SimSun"
    set_explicit_fonts(st._element.get_or_add_rPr(), "宋体", "SimSun")
    st.font.size = Pt(12)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLACK)
    st.paragraph_format.line_spacing = 1.5
    st.paragraph_format.first_line_indent = Pt(0)
    st.paragraph_format.keep_with_next = True

for lname in ["List Bullet", "List Number"]:
    st = styles[lname]
    st.font.name = "SimSun"
    set_explicit_fonts(st._element.get_or_add_rPr(), "宋体", "SimSun")
    st.font.size = Pt(12)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = 1.5

if "Code Block" not in styles:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code Block"]
code_style.font.name = "Consolas"
set_explicit_fonts(code_style._element.get_or_add_rPr(), "等线", "Consolas")
code_style.font.size = Pt(10.5)
code_style.paragraph_format.left_indent = Inches(0.22)
code_style.paragraph_format.right_indent = Inches(0.12)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(6)
code_style.paragraph_format.line_spacing = 1.0
code_style.paragraph_format.first_line_indent = Pt(0)

caption_style = styles["Caption"]
caption_style.font.name = "SimSun"
set_explicit_fonts(caption_style._element.get_or_add_rPr(), "宋体", "SimSun")
caption_style.font.size = Pt(10.5)
caption_style.font.bold = False
caption_style.font.color.rgb = RGBColor.from_string(BLACK)
caption_style.paragraph_format.first_line_indent = Pt(0)


def add_field(paragraph, code, placeholder=""):
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar, instrText, separate, text, end])


def add_p(text="", bold_lead=None, style=None, align=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    if bold_lead and text.startswith(bold_lead):
        font_run(p.add_run(bold_lead), bold=True)
        font_run(p.add_run(text[len(bold_lead):]))
    else:
        font_run(p.add_run(text))
    return p


def bullet(text):
    return add_p(text, style="List Bullet")


def number(text):
    return add_p(text, style="List Number")


def heading(text, level=1):
    return doc.add_heading(text, level=level)


def callout(label, text, fill=PALE, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)
    font_run(p.add_run(label + "  "), bold=True, color=color)
    font_run(p.add_run(text))
    return p


def add_table(headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    font_size = max(font_size, 10.5)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, header_fill)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        font_run(p.add_run(str(h)), size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = t.add_row().cells
        prevent_row_split(t.rows[-1])
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            font_run(p.add_run(str(val)), size=font_size)
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def page_break():
    doc.add_page_break()


def code(text):
    p = doc.add_paragraph(style="Code Block")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FA")
    pPr.append(shd)
    font_run(p.add_run(text), size=10.5, color="243447", east="等线", latin="Consolas")
    return p


# Running header/footer
header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
font_run(hp.add_run("PC6M-10  软件设计及开发文档"), size=10.5, color=MUTED)
font_run(hp.add_run("    |    LPC1765FBD100"), size=10.5, color=MUTED)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font_run(fp.add_run("V1.1    第 "), size=10.5, color=MUTED)
add_field(fp, "PAGE", "1")
font_run(fp.add_run(" 页 / 共 "), size=10.5, color=MUTED)
add_field(fp, "NUMPAGES", "1")
font_run(fp.add_run(" 页"), size=10.5, color=MUTED)

# Cover: editorial_cover
doc.add_paragraph().paragraph_format.space_after = Pt(110)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.first_line_indent = Pt(0)
p.paragraph_format.first_line_indent = Pt(0)
font_run(p.add_run("软件技术文档"), size=12, bold=True, color=GOLD, east="黑体", latin="黑体")
p.paragraph_format.space_after = Pt(18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font_run(p.add_run("PC6M-10 三相 SCR 控制板"), size=22, bold=True, color=NAVY, east="黑体", latin="黑体")
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font_run(p.add_run("软件设计及开发文档"), size=18, bold=True, color=BLUE, east="黑体", latin="黑体")
p.paragraph_format.space_after = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font_run(p.add_run("基于 LPC1765FBD100 的六路移相触发、闭环控制与人机交互软件"), size=12.5, color=MUTED)
p.paragraph_format.space_after = Pt(90)
add_table(["文档版本", "适用固件", "编制日期"], [["V1.1", "PC6M-10 / ST33C", "2026-08-31"]], [2600, 3760, 3000], header_fill=LIGHT_GRAY, font_size=10)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(35)
font_run(p.add_run("内部设计、开发、测试与维护使用"), size=10.5, italic=True, color=MUTED)
page_break()

heading("文档控制", 1)
add_table(["项目", "内容"], [
    ["文档名称", "PC6M-10 三相 SCR 控制板软件设计及开发文档"],
    ["目标平台", "NXP LPC1765FBD100，ARM Cortex-M3"],
    ["文档定位", "软件总体设计、模块设计、接口说明、开发构建、测试、烧写与维护规范"],
    ["适用人员", "嵌入式开发、测试、硬件联调、生产维护及技术支持人员"],
    ["代码基线", "firmware/ 可编译工程及当前 main 分支"],
    ["维护原则", "软件行为、接口或构建流程变化时同步更新本文档"],
], [1900, 7460], font_size=10)

heading("修订记录", 2)
add_table(["版本", "日期", "修订内容"], [
    ["V1.0", "2026-08-31", "首次形成完整软件设计及开发文档"],
    ["V1.1", "2026-08-31", "修正封面及一至三级标题的黑体字体映射"],
], [1200, 1800, 6360], font_size=10)

heading("阅读说明", 2)
add_p("本文档以当前软件源码为设计基线，描述产品背景、系统架构、模块职责、接口协议、控制流程、数据管理以及开发交付方法。涉及强电和门极驱动的测试必须遵守分级放行要求。")
callout("安全提示", "PC6M-10 属于功率控制设备。常规开发、烧写和功能验证应断开市电、SCR 门极及功率负载，仅使用隔离限流控制电。强电带载测试须另行进行风险评估并由具备资质的人员实施。", fill="FFF8E8", color=GOLD)

heading("目录", 1)
toc = doc.add_paragraph()
toc.paragraph_format.first_line_indent = Pt(0)
add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "在 Word 中更新域以刷新目录")
page_break()

heading("1  项目背景", 1)
heading("1.1 产品背景", 2)
add_p("PC6M-10 是面向三相交流功率调节场景的嵌入式控制板。系统通过三路同步输入识别电网相位，依据运行模式、给定值、反馈量和保护状态计算触发时刻，并通过六路门极控制输出驱动外部晶闸管功率单元。控制板同时提供本地显示与按键操作、RS485 远程通信、参数掉电保存、运行/报警继电器以及多类电气保护。")
add_p("软件运行于 NXP LPC1765FBD100 微控制器，采用无操作系统的中断加周期主循环架构。软件既要完成微秒级触发输出，又要承担人机界面、模拟量采样、闭环调节、通信协议和参数管理，因此设计重点是确定性、安全封锁、可维护性以及跨模块状态一致性。")

heading("1.2 建设目标", 2)
for x in [
    "形成结构清晰、可重复构建的 Cortex-M3 裸机固件工程。",
    "稳定产生三相六路 SCR 移相触发脉冲，并在停机或故障时进入安全封锁态。",
    "支持恒压、恒流和开环三种运行方式，以及通讯、本地、定值三种控制来源。",
    "提供 12864 LCD 菜单、编码器/按键、运行停止输入和状态指示。",
    "通过 UART3/隔离 RS485 提供 Modbus-RTU 监控与参数设置。",
    "使用 AT24C02C EEPROM 保存配置、累计量和运行状态。",
    "提供可审计的构建、测试、烧写、回退和分级实机验证流程。",
]: bullet(x)

heading("1.3 设计范围", 2)
add_table(["范围内", "范围外"], [[
    "MCU 固件、LCD、输入扫描、ADC、EEPROM、Modbus、状态机、闭环、触发输出、继电器与指示灯",
    "功率器件选型、门极驱动电路设计、强电柜体设计、上位机应用实现、生产工装机械结构"
]], [4680, 4680], font_size=9.5)

heading("1.4 关键质量属性", 2)
add_table(["属性", "设计要求"], [
    ["安全性", "失去运行条件、发生急停或保护动作时撤销触发；烧写和带载执行分级门禁。"],
    ["确定性", "同步和触发由外部中断与硬件定时器驱动，主循环不承担精确定时。"],
    ["可靠性", "看门狗持续监控；参数采用 live/shadow 机制并支持掉电保存。"],
    ["可维护性", "源码按业务模块拆分，寄存器、常量、全局数据与业务逻辑分层管理。"],
    ["可测试性", "提供静态、主机执行、协议、状态机、闭环和中断序列测试。"],
], [1800, 7560], font_size=9.5)

heading("2  系统概述", 1)
heading("2.1 硬件平台", 2)
add_table(["项目", "配置"], [
    ["主控", "NXP LPC1765FBD100，ARM Cortex-M3，片上 Flash 256 KiB"],
    ["时钟", "12 MHz 晶振输入，软件配置 PLL0/PLL1 和外设时钟树"],
    ["显示", "12864 图形 LCD，8 位并行数据总线和页式显示控制"],
    ["参数存储", "AT24C02C，7 位地址 0x53，GPIO 模拟 I2C"],
    ["通信", "UART3 + ADM2483 隔离 RS485，Modbus-RTU 从站"],
    ["模拟量", "ADC0 六通道轮询采样，覆盖三相电流、输出电流、输出电压和给定测量"],
    ["同步输入", "EINT1/EINT2/EINT3：P2.11/P2.12/P2.13，下降沿触发"],
    ["触发输出", "G1-G6：P0.17/P0.15/P0.18/P2.9/P0.19/P0.16；另保留扩展输出资源"],
    ["继电器", "P0.20 备用、P0.21 报警、P0.22 运行"],
], [1850, 7510], font_size=9.3)

heading("2.2 软件运行模型", 2)
add_p("软件采用前后台结构。前台由 TIMER0、TIMER1、TIMER2、EINT1-3、UART3 和 WDT 中断组成，负责周期节拍、同步捕获、触发脉冲、串口收发和故障监控；后台主循环在 TIMER0 节拍到达后依次执行采样、输入扫描、状态机、输出计算、看门狗喂狗、串口超时处理和协议分发。")
code("TIMER0 tick\n  -> ADC 扫描\n  -> 输入事件扫描\n  -> ADC 补充扫描\n  -> 人机/运行状态机\n  -> 输出与保护计算\n  -> 看门狗喂狗\n  -> UART 接收超时管理\n  -> Modbus 帧分发")
callout("设计要点", "精确触发时序由 TIMER1/TIMER2 和外部中断闭环完成；LCD、参数保存、通信解析等非硬实时工作留在周期主循环中。", fill=LIGHT_BLUE)

heading("2.3 分层架构", 2)
add_table(["层次", "组成", "职责"], [
    ["应用与交互层", "状态机、菜单、LCD、Modbus 分发", "运行控制、参数编辑、信息呈现和远程操作"],
    ["控制层", "输出级、闭环、保护、软起停", "计算给定、触发角、保护状态和运行状态"],
    ["服务层", "参数系统、CRC、输入消抖", "配置持久化、协议校验、输入事件整形"],
    ["驱动层", "GPIO、ADC、UART3、I2C、TIMER、EINT、WDT", "直接访问 LPC1765 外设"],
    ["平台层", "startup.s、lpc1765.ld、data_image.s", "启动、向量、内存布局和镜像初始化"],
], [1500, 2860, 5000], font_size=9.2)

heading("3  软件总体设计", 1)
heading("3.1 启动流程", 2)
add_table(["顺序", "动作", "说明"], [
    ["1", "Reset_Handler", "复制 SRAM0 初始化镜像，清零 SRAM0 工作区；初始化 SRAM1 的 data/bss。"],
    ["2", "SystemInit", "配置振荡器、PLL、电源和外设时钟。"],
    ["3", "GPIO/LCD/输入", "配置引脚复用、LCD、输入方向和初始电平。"],
    ["4", "TIMER0/I2C/ADC", "建立系统节拍、参数总线和模拟量采样。"],
    ["5", "load_config", "从 EEPROM 装载配置并生成运行态派生值。"],
    ["6", "输出与中断", "初始化 TIMER1/TIMER2、EINT1-3、UART3 和输出安全态。"],
    ["7", "显示与看门狗", "显示启动界面并启用 WDT。"],
    ["8", "main loop", "进入按 TIMER0 节拍驱动的永久循环。"],
], [900, 2300, 6160], font_size=9.3)

heading("3.2 内存布局", 2)
add_table(["区域", "地址/容量", "用途"], [
    ["FLASH", "0x00000000 / 256 KiB", "向量表、CRP 控制字、代码、只读常量和 SRAM 初始化镜像"],
    ["SRAM0", "0x10000000 / 32 KiB", "业务数据镜像、运行变量、BSS 和主栈"],
    ["SRAM1", "0x2007C000 / 16 KiB", "编译期全局符号和工程辅助数据"],
    ["初始栈顶", "0x100029C8", "Reset 向量设置的 MSP 初值"],
    ["CRP 字", "Flash 0x000002FC", "固定为 0xFFFFFFFF，保持无代码读保护"],
], [1500, 2400, 5460], font_size=9.5)
add_p("链接脚本将 CRP 字固定在 0x2FC，并使普通代码避开该位置。构建二进制使用 0xFF 填充空洞，保持与 Flash 擦除态一致。向量表前八个字包含 LPC 启动校验和。")

heading("3.3 中断体系", 2)
add_table(["中断", "来源", "主要职责"], [
    ["WDT", "看门狗超时", "清除超时标志并累计异常次数"],
    ["TIMER0", "周期匹配", "置主循环 tick、维护相位计数和倒计时"],
    ["TIMER1", "高速扫描", "按 240 步序列生成六路门极脉冲窗口"],
    ["TIMER2", "单次延时", "根据触发角延时后启动触发扫描"],
    ["EINT1/2/3", "三相同步下降沿", "识别相位事件、更新模式并启动 TIMER2"],
    ["UART3", "RDA/THRE", "接收组帧和发送缓冲推进"],
], [1400, 2000, 5960], font_size=9.4)
add_p("工程未主动配置 NVIC 优先级，使用芯片复位后的默认优先级关系。新增中断或改变实时路径前，应评估 TIMER1 触发抖动、最长关中断时间和 ISR 嵌套影响。")

heading("3.4 安全状态", 2)
add_p("触发输出具有硬件极性约束。软件的统一封锁函数将全部触发相关 GPIO 置为高电平，使输出级不产生有效门极脉冲。停机、急停、故障、复位和上电初始化均应能够到达该状态。继电器和状态指示由独立封装函数控制，避免业务代码直接散落写 GPIO。")

heading("4  模块详细设计", 1)
modules = [
    ("01_startup.c", "系统启动与调度", "时钟、WDT、TIMER0、main 初始化序列和周期主循环"),
    ("02_lcd_display.c", "显示驱动", "LCD GPIO、命令/数据总线、字符、数字、反显和固定页面"),
    ("03_input_debounce.c", "输入与消抖", "编码器、RUN/STOP、急停、复位、外故障和联锁输入"),
    ("04_i2c.c", "EEPROM 总线", "GPIO 模拟 I2C 的起停、字节收发和寄存器读写"),
    ("05_adc.c", "模拟量采样", "ADC0 初始化、启动、DONE 等待和六通道均值换算"),
    ("06_param_system.c", "参数管理", "EEPROM 装载、默认初始化、live/shadow 同步和派生值"),
    ("07_state_machine.c", "人机与运行状态机", "菜单导航、编辑、显示刷新、运行/停止、故障与复位"),
    ("08_uart3_modbus.c", "通信底层", "UART3、接收帧、发送、CRC16、寄存器读写"),
    ("08_modbus_dispatch.c", "协议分发", "0x03/0x06/0x10 校验、异常响应和业务分派"),
    ("09_output_stage.c", "触发与保护核心", "引脚、TIMER/EINT、输出计算、软起停和门极脉冲"),
    ("10_relay_led.c", "继电器与指示灯", "运行、报警、备用继电器和状态 LED 电平封装"),
    ("11_auth.c", "设备识别服务", "GPIO 串行挑战应答与重试；当前版本主流程固定放行"),
    ("12_closed_loop.c", "闭环控制", "误差分区、PID 计算、积分累加、限幅和节流包装"),
    ("13_gpio_init.c", "NVIC/GPIO2", "中断使能和 GPIO2 串行接口初始化"),
]
add_table(["源文件", "模块", "职责"], modules, [2200, 2050, 5110], font_size=8.8)

heading("4.1 启动与系统调度", 2)
add_p("`main()` 先完成平台和业务外设初始化，再进入永久循环。循环只在 `tick_ready==1` 时运行一次，执行结束前必须调用 `wd_feed()`。如果 P0.2/P0.3 安全联锁在启动时处于异常状态，系统进入联锁页面并只保留输入扫描、频率同步、停机预设和看门狗服务。")
add_p("看门狗配置采用 WDFEED 0xAA、0x55 双字节序列。TIMER0 初始化先通过 TCR=2 复位计数器，再设置 PR、MR0、MCR 和中断。任何修改都必须保持寄存器写入顺序。")

heading("4.2 LCD 显示", 2)
add_p("LCD 驱动采用 8 位并行数据输出。底层通过 `lcd_data_byte()` 设置数据位，通过 RS 和 E 控制线区分命令与数据；上层支持 8×16 ASCII、16×16 中文、反显、高亮、定点小数、三位/四位/五位数字和带符号角度显示。字符串通过统一映射接口访问，以保证显示内容与代码地址解耦。")
bullet("`disp_clear()` 分别清除控制器上下两半的 8 个页面。")
bullet("`disp_string()` 根据字节编码选择 ASCII 或双字节中文字符。")
bullet("`lcd_ctrl_line()` 控制背光；状态机在长时间无操作后关闭背光，按键事件恢复。")

heading("4.3 输入扫描与事件", 2)
add_table(["输入", "引脚", "处理"], [
    ["编码器/菜单键", "P1.16/P1.17、P0.27/P0.28、P2.24/P2.25", "组合采样、方向锁存、短按/快进/快退事件"],
    ["RUN/STOP", "P0.28/P0.27", "阈值 50；支持点动和保持两种启动方式"],
    ["交流检测", "P0.9", "阈值 15 的高电平确认"],
    ["外部故障", "P1.16", "高低双向消抖，阈值 250"],
    ["复位", "P1.17", "高低双向消抖，阈值 50"],
    ["急停", "P0.6", "高低双向消抖，阈值 50"],
    ["启动联锁", "P0.2/P0.3", "上电检查和运行状态记录"],
], [1650, 2600, 5110], font_size=9.0)
add_p("所有 FIO 地址偏移均按字节计算。不得对 `uint32_t*` 直接增加寄存器字节偏移，否则会产生四倍地址偏差。")

heading("4.4 EEPROM 与参数系统", 2)
add_p("参数存储器使用 AT24C02C，地址 0x53。底层总线由 P0.10(SDA) 和 P0.11(SCL) 模拟，写事务后保留 EEPROM 内部写周期延时。参数系统维护 live 值和 shadow 值：菜单或通信修改 live 后，只有检测到与 shadow 不同才写 EEPROM，并同步 shadow，降低不必要的擦写次数。")
add_p("配置按两个 magic 标记区组织。上电时标记有效则读取参数；标记无效则把当前默认配置写入 EEPROM。16 位参数按高低字节保存。运行时间、控制模式、保护参数、通信参数、PID 参数和输出标定参数均纳入持久化。")
callout("开发约束", "增加持久参数时必须同时修改装载路径、live/shadow 定义、保存路径、默认值、菜单或 Modbus 入口，并明确字节序与写入范围。", fill=LIGHT_BLUE)

heading("4.5 ADC 数据采集", 2)
add_table(["通道", "信号", "用途"], [
    ["AD0.0", "IC", "C 相电流"], ["AD0.1", "IB", "B 相电流"], ["AD0.2", "IA", "A 相电流"],
    ["AD0.3", "IF", "输出电流反馈与保护"], ["AD0.4", "Uf", "输出电压反馈与保护"],
    ["AD0.5", "Ug", "给定/输入测量"],
], [1300, 2200, 5860], font_size=9.5)
add_p("采样流程为选择通道、启动转换、轮询 AD0GDR 的 DONE 位、提取 12 位结果，再按通道执行累计、平均和量程换算。AD0GDR 位于 ADC0 基址 +0x04，访问代码必须使用字节地址偏移。")

heading("4.6 状态机与菜单", 2)
add_p("`state_machine(key)` 是本地交互和运行状态的中心。MENU 表示一级页面，MENU2 表示当前项目，MENU3 表示查看或编辑状态；TIMEOUT/TIMEOUT3 分别负责页面超时与闪烁刷新。状态机同时读取运行、故障、急停、复位、控制来源和启动方式等全局状态。")
add_table(["页面组", "项目数", "主要内容"], [
    ["主页面", "1", "输入信号、输出反馈、运行状态、故障和本地给定"],
    ["基本参数", "16", "运行模式、量程、限制、软起停、控制方式、启动方式和起始相位"],
    ["保护参数", "10", "过欠压、IF/CT 过载、缺相和三相平衡"],
    ["通讯参数", "4", "从站地址、波特率、校验/帧格式和通信检测"],
    ["PID 参数", "9", "PID 档位、P/I 和闭环误差分区参数"],
    ["维护页面", "若干", "相位校准、恢复设置、版本信息和运行记录"],
], [1800, 1100, 6460], font_size=9.2)
add_p("编辑态通过周期性整页重绘与当前值擦除产生闪烁。枚举参数的上限行为不完全一致：运行模式允许环绕，而控制方式、启动方式、急停、反馈和输入选择在上限处钳位。维护时不得把所有枚举编辑逻辑合并成一个未经验证的通用函数。")

heading("4.7 UART3 与 Modbus-RTU", 2)
add_p("UART3 通过 ADM2483 接入隔离 RS485。UART 中断处理接收数据可用和发送保持寄存器空两类事件；接收侧维护帧缓冲和超时，主循环在帧完整后执行协议分发。CRC16 使用多项式 0xA001、初值 0xFFFF，报文中的 CRC 低字节在前。")
add_table(["功能码", "功能", "主要校验"], [
    ["0x03", "读保持寄存器", "站址、CRC、起始寄存器 1..63、数量和响应长度"],
    ["0x06", "写单寄存器", "站址、CRC、寄存器和值范围；成功后同步参数"],
    ["0x10", "写多个寄存器", "站址、CRC、数量、字节数和连续地址"],
], [1300, 2600, 5460], font_size=9.4)
add_p("Modbus 16 位地址、数量和值采用大端拼接；CRC 字段仍按 Modbus 规定低字节在前。设备只响应与当前从站地址完全匹配的报文，不实现地址 0 广播。")

heading("4.8 输出级与触发时序", 2)
add_p("输出级把运行模式、给定、反馈、限制、故障和软起停状态合成为触发角。TIMER2 负责从同步沿到触发窗口的单次延时；TIMER1 进入约 240 步的高速扫描，在六个 60° 窗口内产生门极脉冲列。50 Hz 与 60 Hz 使用不同的 TIMER1 匹配参数。")
add_table(["资源", "职责", "关键设计"], [
    ["EINT1-3", "三相同步捕获", "下降沿触发，更新同步模式和相序状态"],
    ["TIMER2", "相位延时", "MR0 由目标触发角计算，匹配后启动 TIMER1"],
    ["TIMER1", "脉冲序列", "按计数窗口切换 G1-G6，周期结束停止或等待下一同步"],
    ["gpio_outputs_set", "统一封锁", "全部门极相关输出置高，作为停机/故障安全态"],
], [1600, 2500, 5260], font_size=9.2)
add_p("触发角使用 180° 基准和角度比例系数换算为定时器值。给定变化应使触发时刻单调变化。任何涉及 TIMER1 ISR、GPIO 掩码、窗口边界或电平极性的修改，都必须重新进行六路空载波形验证。")

heading("4.9 闭环控制", 2)
add_p("闭环模块支持恒压和恒流反馈。算法先根据误差幅值和配置阈值选择分段除数，再按带符号误差执行比例、积分和差分项组合，输出累加到积分状态并执行上下限钳位。包装函数通过重算计数控制 PID 更新频率，未到重算周期时返回上一次缓存结果。")
bullet("误差及其中间量必须保持有符号 32 位语义。")
bullet("分段阈值和增益来自持久参数，活动参数组在运行模式变化时刷新。")
bullet("软启动、软停止、开环和闭环共用输出限幅与安全封锁条件。")

heading("4.10 继电器、指示灯与设备识别", 2)
add_p("继电器和状态 LED 均通过单一职责函数设置 SET/CLR 位，业务层不直接操作对应 GPIO。运行、报警和备用继电器的动作由状态机及输出级决定。设备识别模块保留 GPIO 串行挑战应答和重试能力；当前软件版本在主流程中固定允许运行，因此该模块不作为安全边界。")

heading("5  数据与参数设计", 1)
heading("5.1 全局数据分类", 2)
add_table(["分类", "示例", "管理原则"], [
    ["实时采样", "IA/IB/IC、IF、Uf、Ug", "由 ADC 模块更新，状态机、保护和通信只读或派生使用"],
    ["用户配置", "模式、量程、限制、软起停、通信、PID", "live/shadow + EEPROM 持久化"],
    ["运行状态", "RUN、FAULT、STOP_PEND、触发模式", "主循环与 ISR 共享，使用 volatile"],
    ["界面状态", "MENU、MENU2、MENU3、TIMEOUT", "由状态机维护，显示函数仅消费"],
    ["通信状态", "RX/TX 缓冲、长度、帧状态、从站地址", "UART ISR 与主循环分工，避免并发覆盖"],
    ["控制状态", "PID 误差、积分、缓存输出、软起停计数", "闭环和输出级专有，重启时按初始化策略恢复"],
], [1700, 2900, 4760], font_size=9.0)

heading("5.2 基本参数", 2)
add_table(["参数", "范围/选项", "作用"], [
    ["运行模式", "0..2：恒压/恒流/开环", "选择反馈和输出算法"],
    ["电压/电流量程", "0..6000", "显示、反馈换算和保护上限基准"],
    ["互感器比", "0..6000", "三相电流换算和 CT 过载保护"],
    ["电压/电流限制", "不超过对应量程", "输出限制与保护"],
    ["软起/软停时间", "0..200", "控制斜坡速度"],
    ["相位限制", "0..180", "触发角限制"],
    ["控制方式", "0 通讯 / 1 本地 / 2 定值", "选择给定来源"],
    ["启动方式", "0 点动 / 1 自锁", "定义 RUN/STOP 输入行为"],
    ["起始相位", "0..180", "输出下限和启动角"],
], [2200, 2600, 4560], font_size=9.2)

heading("5.3 保护参数", 2)
add_table(["保护", "阈值", "延时/设置"], [
    ["过压", "不超过电压量程", "0..200"], ["欠压", "不超过电压量程", "0..200"],
    ["IF 过载", "不超过电流量程", "0..200"], ["CT 过载", "不超过互感器比", "0..200"],
    ["缺相", "关/开", "同步缺失判定"], ["三相平衡", "0..60%", "三相电流不平衡阈值"],
], [2300, 2800, 4260], font_size=9.2)

heading("5.4 参数修改一致性", 2)
number("本地菜单进入编辑态后只修改 live 参数，并立即应用范围钳位。")
number("确认或超时退出时调用参数同步服务。")
number("Modbus 0x06 写入通过寄存器专属校验后更新 live，并调用同步服务。")
number("同步服务比较 live 与 shadow，只写入发生变化的 EEPROM 字节。")
number("需要重启生效的通信参数应在界面或上位机侧明确提示。")

heading("6  外部接口设计", 1)
heading("6.1 GPIO 接口", 2)
add_table(["功能", "引脚", "方向/说明"], [
    ["G1-G6", "P0.17/P0.15/P0.18/P2.9/P0.19/P0.16", "输出；六路门极触发"],
    ["同步输入", "P2.11/P2.12/P2.13", "输入；EINT1-3 下降沿"],
    ["EEPROM", "P0.10/P0.11", "双向 SDA / 输出 SCL"],
    ["RS485 DE/RE", "P1.29", "输出；与 SWCLK 复用"],
    ["电压反馈", "P1.30 / AD0.4", "模拟输入；与 SWDIO 复用"],
    ["继电器", "P0.20/P0.21/P0.22", "备用/报警/运行"],
    ["状态 LED", "P1.20..P1.23", "模式与状态指示"],
    ["ISP/SWD", "P2.10、P1.30、P1.29、nRESET", "启动恢复和开发接口"],
], [1800, 3600, 3960], font_size=9.0)
callout("调试限制", "P1.29 和 P1.30 在应用启动后分别用于 RS485 控制和 ADC，因此运行态 SWD 可能失联。调试连接优先使用 100 kHz；需要时采用 connect-under-reset，或通过 ISP 恢复。", fill="FFF8E8", color=GOLD)

heading("6.2 Modbus 常用寄存器", 2)
add_table(["寄存器", "方向", "含义", "说明"], [
    ["40", "读", "Ug 给定测量", "AD0.5 实时测量"],
    ["40", "写", "远程给定", "值不大于 1000；读写语义不同"],
    ["41/42/43", "读", "IA/IB/IC", "三相电流"],
    ["44", "读", "IF", "输出电流反馈"],
    ["45", "读", "Uf", "输出电压反馈"],
    ["61", "写", "远程输出使能", "关联 P0.20 备用继电器，写测试需安全授权"],
    ["62", "读写", "起始相位/输出下限", "范围 0..180"],
], [1300, 1200, 2700, 4160], font_size=9.0)
add_p("寄存器 24/25、40 等存在业务上的读写不对称。上位机设计不能假设所有寄存器写后回读等于写入值，应按接口定义分别处理读值和写值。")

heading("6.3 LCD 与本地操作", 2)
add_p("本地界面以四行显示为主要布局，通过反显标识当前项目，通过值的周期显示/擦除表示编辑状态。主页面提供输入信号、输出反馈、运行/停止和故障信息。编码器或 UP/DOWN 产生增减事件，SET 在查看、编辑和下级页面之间切换，超时自动保存并返回主页面。")

heading("7  控制流程设计", 1)
heading("7.1 运行状态转换", 2)
code("上电初始化\n  -> 安全联锁检查\n  -> 停机待命\n  -> 收到 RUN 且无故障\n  -> 软启动\n  -> 稳态运行（开环/恒压/恒流）\n  -> STOP 或故障\n  -> 软停止或立即封锁\n  -> 停机待命")
add_p("急停、外部故障、过压、过流、缺相和相序异常具有比普通停止更高的安全优先级。进入故障处理后应清除 RUN、设置停机请求并封锁门极；复位只有在对应条件释放后才允许恢复。")

heading("7.2 控制来源", 2)
add_table(["控制方式", "给定来源", "本地按键行为"], [
    ["通讯", "Modbus 远程给定", "主页面不直接修改给定"],
    ["本地", "现场输入/测量链路", "主页面不直接修改给定"],
    ["定值", "保存的 MANUAL 定值", "主页面 UP/DOWN 每次调整 0.1%，支持长按"],
], [1600, 3100, 4660], font_size=9.4)

heading("7.3 软起动与软停止", 2)
add_p("软起动从安全初值逐步移动到目标输出，软停止按配置时间降低输出并在结束后封锁门极。斜坡状态与运行/停止请求共同决定输出级状态码。修改斜坡算法时必须覆盖时间为 0、最小值、最大值、中途故障和中途停止场景。")

heading("7.4 故障与保护", 2)
add_p("保护计算综合 Uf、IF、三相电流、同步相位、外部故障和急停输入。部分保护带时间计数，避免短时扰动造成误动作。故障状态同时影响 LCD、报警继电器、运行继电器和触发封锁。故障复位流程会清理运行请求并显示复位/重启提示。")

heading("8  工程结构与编码规范", 1)
heading("8.1 工程目录", 2)
code("firmware/\n  startup.s             向量表与复位入口\n  lpc1765.ld            Flash/SRAM/CRP 链接布局\n  data_image.s          SRAM 初始数据镜像\n  inc/firmware_state.h  运行时状态与外设的语义地址映射\n  inc/firmware_parameters.h 参数区与 EEPROM 镜像的语义地址映射\n  inc/firmware_api.h    模块公共接口\n  src/                  业务与驱动模块\n  build.sh              推荐构建入口\n  Makefile              make 构建入口\n  firmware.elf/.bin/.hex/.map  构建产物\ntest/                   自动化测试\ntools/jlink/            仓库版 J-Link 工具\ntools/w8/               串口、Modbus、波形和实机辅助工具")

heading("8.2 编码规范", 2)
for x in [
    "外设寄存器和 ISR/主循环共享状态必须使用 volatile。",
    "寄存器偏移使用 `uintptr_t` 或整数地址表达字节偏移，不对有类型指针直接加字节数。",
    "访问宽度必须与硬件寄存器和数据定义一致；8 位参数不得用 32 位写入。",
    "状态机修改应保留原有 fall-through、共享尾部、超时和多周期行为。",
    "协议字段显式表达字节序；Modbus 16 位数据高字节在前、CRC 低字节在前。",
    "定时、GPIO 和喂狗序列不得因重构而改变写入顺序。",
    "常量集中放入 `inc/consts.h`；寄存器定义集中放入 `inc/reg.h`。",
    "新增全局状态应明确所有者、写入者、读取者、位宽、初始值和持久化要求。",
]: bullet(x)

heading("8.3 变更影响分析", 2)
add_table(["变更类型", "至少检查"], [
    ["GPIO/引脚", "PINSEL、DIR、SET/CLR 极性、安全态、板级接线、空载波形"],
    ["定时器/ISR", "频率、窗口边界、清中断、最长执行时间、连续中断"],
    ["参数", "默认值、范围、live/shadow、EEPROM、菜单、Modbus、文档"],
    ["状态机", "菜单导航、故障态、多 tick、显示参数、超时保存"],
    ["Modbus", "0x03/0x06/0x10、异常码、字节序、CRC、寄存器兼容性"],
    ["闭环", "正负误差、死区、分段边界、上下限、软起停和带载风险"],
], [1900, 7460], font_size=9.3)

heading("9  构建与开发环境", 1)
heading("9.1 工具依赖", 2)
add_table(["工具", "版本/用途"], [
    ["Git for Windows", "版本管理，并提供 Git Bash"],
    ["Arm GNU Toolchain", "14.2.Rel1；arm-none-eabi-gcc/objcopy/size"],
    ["Python", "3.12；测试、串口、Modbus 和辅助脚本"],
    ["J-Link", "使用仓库 tools/jlink/JLink.exe 及配套驱动"],
    ["Flash Magic", "ISP 恢复通道和 UART0 设备操作"],
], [2200, 7160], font_size=9.5)

heading("9.2 编译参数", 2)
code("-mcpu=cortex-m3 -mthumb -mfloat-abi=soft\n-Os -ffreestanding -fno-builtin -Wall\n-T lpc1765.ld -nostdlib -lgcc\n-Wl,--gc-sections -Wl,-Map,firmware.map")
add_p("推荐从 Git Bash 执行 `build.sh`。脚本优先查找 14.2.Rel1，再查找常见安装目录和 PATH。构建前自动清理旧对象，逐模块编译，链接后生成 ELF、HEX、BIN 和 MAP。")

heading("9.3 标准构建", 2)
code("cd firmware\nbash build.sh")
add_table(["产物", "用途"], [
    ["firmware.elf", "带符号可执行文件，用于符号调试、地址定位和测试"],
    ["firmware.bin", "从 0x00000000 起烧写的二进制镜像"],
    ["firmware.hex", "Intel HEX，适用于部分 ISP/生产工具"],
    ["firmware.map", "链接映射、符号地址和段布局审计"],
], [2200, 7160], font_size=9.5)

heading("9.4 构建后检查", 2)
for x in [
    "编译和链接无错误；警告数量无意外增加。",
    "`arm-none-eabi-size firmware.elf` 的 text/data/bss 变化符合预期。",
    "firmware.bin 不超过 256 KiB。",
    "Flash 0x2FC 为 0xFFFFFFFF，向量表校验和有效。",
    "MAP 中关键 ISR、main 和核心模块符号存在，业务函数没有异常消失。",
    "记录 firmware.bin SHA-256，作为烧写与测试批次标识。",
]: bullet(x)

heading("10  测试与质量保证", 1)
heading("10.1 测试分层", 2)
add_table(["层级", "覆盖重点", "通过要求"], [
    ["静态检查", "CRC 表、Modbus 映射、参数同步结构、访问宽度和符号覆盖", "全部检查通过"],
    ["主机执行测试", "CRC、协议分发、参数同步、闭环计算", "边界和错误路径通过"],
    ["状态序列测试", "输入、状态机、多 tick、ISR、继电器和输出级", "状态末态和外设写序符合设计"],
    ["控制电实测", "启动、LCD、按键、EEPROM、看门狗、继电器安全态", "无异常复位或危险动作"],
    ["空载波形", "三相同步、六路脉冲、相序、触发角、封锁", "全部通道符合时序判据"],
    ["标定/低压", "Modbus、ADC 双点标定、闭环、软起停和保护", "误差与保护指标达标"],
], [1700, 4740, 2920], font_size=8.9)

heading("10.2 自动化测试入口", 2)
code("python test/run_tests.py\npython test/run_tests.py crc16")
add_p("修改模块后先运行对应测试，再运行完整测试入口。测试应覆盖正常值、零值、上下边界、边界外一位、连续调用和错误输入。针对历史故障必须保留可稳定触发问题的定向回归用例。")

heading("10.3 重点测试矩阵", 2)
add_table(["对象", "建议维度"], [
    ["输入消抖", "引脚高低 × 计数器边界 × 锁存状态 × 单次/保持模式"],
    ["状态机", "MENU/MENU2/MENU3 × key × FAULT × RUN × 多 tick"],
    ["Modbus", "功能码 × 地址 × 数量 × 值边界 × CRC/长度/站址"],
    ["闭环", "正负零误差 × 死区 × 分段阈值 × 上下钳位 × 通道"],
    ["输出级", "运行模式 × 控制方式 × 故障 × 软起停 × 触发相位"],
    ["中断", "模式 × 计数回绕 × 连续调用 × 清标志和 GPIO 写序"],
], [1800, 7560], font_size=9.3)

heading("10.4 实机分级放行", 2)
add_table(["阶段", "内容", "放行条件"], [
    ["0", "上电前检查和恢复准备", "固件与接线确认，功率级/门极/负载断开"],
    ["1", "控制电冒烟", "启动、LCD、按键、EEPROM、WDT、继电器无异常"],
    ["2", "三相同步和空载触发", "六路脉冲、相序、触发角、停机/保护撤销全部通过"],
    ["3", "Modbus 和 ADC 标定", "只读通信、双点标定和掉电保存通过"],
    ["4", "低压限流到带载", "前述全部通过并完成风险评估与作业许可"],
], [1000, 3700, 4660], font_size=9.1)

heading("11  烧写、启动与恢复", 1)
heading("11.1 SWD 主通道", 2)
add_p("SWD 使用仓库打包的 J-Link。连接从 100 kHz 开始，首先确认 VTref、GND、SWDIO 和 SWCLK 接触可靠。应用运行后复用调试脚时，可连接 nRESET 并采用 connect-under-reset。")
code(".\\tools\\jlink\\JLink.exe -CommanderScript tools\\jlink\\check.jlink\n.\\tools\\jlink\\JLink.exe -CommanderScript tools\\jlink\\flash.jlink")
add_p("标准烧写脚本顺序为：连接、保存板内镜像、检查 0x2FC、擦除、写入 firmware.bin、verifybin、复位运行。只有出现 `Verify successful` 才能判定写入成功。")

heading("11.2 ISP 恢复通道", 2)
add_p("当 SWD 因引脚复用或保护状态无法连接时，可使用 UART0 ISP。USB-TTL 的 TX/RX 交叉并共地，复位时把 P2.10 拉低进入 Boot ROM。PC6M 的串口网络可能受到隔离 RS485 器件输出影响，必要时禁用其接收输出。ISP 适合器件识别、读取小范围和擦除恢复；完整镜像写入优先回到 SWD。")

heading("11.3 烧写检查清单", 2)
for x in [
    "断开市电、功率负载和 SCR 门极，仅接控制电。",
    "确认待烧写 firmware.bin 的路径、大小和 SHA-256。",
    "保存板内当前镜像并检查 CRP 状态。",
    "执行 erase/loadbin/verifybin，不跳过校验。",
    "检查 0x2FC 仍为 0xFFFFFFFF。",
    "物理断电再上电，确认固件正常启动。",
    "记录提交号、固件哈希、板号、操作者和日期。",
]: bullet(x)

heading("11.4 回退策略", 2)
add_p("每次现场升级前保留当前可运行镜像和构建信息。新版本控制电验证失败时立即恢复上一版本；SWD 无法连接时使用 ISP 进入 Boot ROM，必要时擦除后再用 SWD 写回。任何回退都必须重新执行 verifybin 和控制电冒烟检查。")

heading("12  调试与故障定位", 1)
heading("12.1 调试方法", 2)
add_table(["现象", "优先检查"], [
    ["J-Link 无法连接", "四根主线接触、VTref、100 kHz、应用是否复用 SWD、nRESET"],
    ["LCD 闪屏/周期复位", "看门狗计数、ADC DONE 等待、HardFault、主循环是否喂狗"],
    ["按键全部无效", "FIO 地址偏移、输入方向、事件计数器和 TIMER0 tick"],
    ["Modbus 无响应", "从站地址、波特率/校验、RS485 DE、帧 CRC、UART3 ISR"],
    ["写参数不保存", "live/shadow 是否变化、EEPROM ACK、同步调用、写周期"],
    ["触发脉冲缺失", "EINT 相位输入、TIMER2 MR0、TIMER1 状态、封锁条件和 GPIO 极性"],
    ["闭环方向异常", "给定/反馈单位、有符号误差、增益组、分段除数和钳位"],
], [2500, 6860], font_size=9.1)

heading("12.2 日志与记录", 2)
add_p("每次问题定位应记录软件提交、固件哈希、板号、供电、接线、参数、输入条件、现象、预期、复现步骤和附件路径。协议问题保存原始十六进制帧；波形问题保存示波器 CSV 和截图；参数问题同时记录菜单值、Modbus 值和掉电重启结果。")

heading("12.3 高风险开发陷阱", 2)
add_table(["陷阱", "防范措施"], [
    ["有类型指针导致偏移缩放", "地址偏移统一转为 uintptr_t 后按字节计算"],
    ["8 位参数被 32 位访问", "按定义使用 uint8_t/uint16_t/uint32_t，关注非对齐地址"],
    ["volatile 丢失", "外设和共享状态在所有强转后仍保持 volatile"],
    ["状态机提前 return", "覆盖共享尾部和多 tick 调用序列"],
    ["枚举环绕/钳位混用", "逐参数定义边界语义并建立边界测试"],
    ["协议字节序混淆", "数据字段和 CRC 字段分别编码，使用已知帧测试"],
    ["最终 GPIO 相同但中间写序错误", "触发、继电器和总线驱动必须检查完整写入序列"],
    ["测试种子被默认值覆盖", "通用初始化后再写目标用例值，并在执行前断言"],
], [2800, 6560], font_size=9.1)

heading("13  发布与维护", 1)
heading("13.1 发布流程", 2)
number("冻结需求、参数表和硬件版本，确认变更范围。")
number("完成代码评审、模块测试、完整回归和构建后检查。")
number("生成 ELF/HEX/BIN/MAP，记录 Git 提交和 firmware.bin SHA-256。")
number("按控制电、空载、通信标定、低压限流逐级验证。")
number("更新软件设计文档、操作文档、参数/协议文档和发布记录。")
number("归档可恢复镜像、测试证据和现场回退步骤。")

heading("13.2 发布检查表", 2)
for x in [
    "源码和文档处于同一提交，工作区无无关改动。",
    "工具链版本为 14.2.Rel1，构建命令和输出已保存。",
    "自动化测试全部通过，失败豁免有书面依据。",
    "CRP、向量、段布局、尺寸和哈希检查通过。",
    "Modbus 兼容性和 EEPROM 迁移影响已评估。",
    "六路触发输出、安全封锁、继电器和保护路径已验证。",
    "现场升级和回退镜像可用，操作人员已明确停止线。",
]: bullet(x)

heading("13.3 扩展开发建议", 2)
add_p("新增功能优先在现有模块边界内扩展。新增通信寄存器应避免改变 1..63 的既有语义；新增参数应预留 EEPROM 地址并考虑旧配置兼容；新增触发模式应使用独立状态和定时器矩阵验证；新增硬件引脚前应检查 PINSEL、SWD/ISP、ADC 和总线复用冲突。")

heading("14  附录", 1)
heading("14.1 常用命令", 2)
code("# 构建\ncd firmware && bash build.sh\n\n# 完整测试\ncd ..\npython test/run_tests.py\n\n# SWD 探测\n.\\tools\\jlink\\JLink.exe -CommanderScript tools\\jlink\\check.jlink\n\n# SWD 烧写\n.\\tools\\jlink\\JLink.exe -CommanderScript tools\\jlink\\flash.jlink")

heading("14.2 关键常量", 2)
add_table(["常量", "值", "意义"], [
    ["ANGLE_FULL", "180", "触发角满量程"],
    ["ANGLE_SCALE", "6333", "角度到定时器值的比例系数（×100）"],
    ["TRIG_PERIOD", "0x2C88", "触发周期/扫描基量"],
    ["TRIG_WINDOW", "0x36", "触发窗口匹配值"],
    ["MR0_50HZ", "0x488", "50 Hz TIMER1 匹配参数"],
    ["MR0_60HZ", "0x261", "60 Hz TIMER1 匹配参数"],
    ["EEPROM_MAGIC_U", "0x55", "配置区域 A 标记"],
    ["EEPROM_MAGIC_f", "0x66", "配置区域 B 标记"],
], [2200, 1800, 5360], font_size=9.3)

heading("14.3 安全停止线", 2)
callout("停止线 1", "未确认接线、固件哈希、CRP 和可恢复镜像，不执行擦除或升级。", fill="FDECEC", color=RED)
callout("停止线 2", "空载六路触发时序或保护撤销未通过，不进入低压和带载测试。", fill="FDECEC", color=RED)
callout("停止线 3", "强电测试必须使用合适的隔离/差分测量手段，并由具备资质的人员执行。", fill="FDECEC", color=RED)

# document settings
settings = doc.settings.element
update = OxmlElement("w:updateFields")
update.set(qn("w:val"), "true")
settings.append(update)
doc.core_properties.title = "PC6M-10 三相 SCR 控制板软件设计及开发文档"
doc.core_properties.subject = "LPC1765FBD100 固件软件设计、开发、构建、测试与烧写"
doc.core_properties.author = "PC6M-10 项目组"
doc.core_properties.keywords = "PC6M-10, LPC1765, SCR, Modbus, firmware"

doc.save(OUT)


def validate_output(path):
    """Fail fast when generation drifts from the maintained document baseline."""
    generated = Document(path)
    full_text = "\n".join(
        [paragraph.text for paragraph in generated.paragraphs]
        + [
            cell.text
            for table in generated.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    forbidden = (
        "逆向", "反编译", "反汇编", "Ghidra", "Unicorn", "A/B", "TODO", "TBD"
    )
    found = [token for token in forbidden if token in full_text]
    if found:
        raise RuntimeError(f"文档包含禁止或占位表述: {found}")

    headings = [
        paragraph for paragraph in generated.paragraphs
        if paragraph.style is not None and paragraph.style.name.startswith("Heading")
    ]
    if len(generated.tables) != 31 or len(headings) != 74:
        raise RuntimeError(
            f"结构数量异常: tables={len(generated.tables)}, headings={len(headings)}"
        )

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        rfonts = generated.styles[style_name]._element.rPr.rFonts
        if rfonts is None or rfonts.get(qn("w:eastAsia")) != "黑体":
            raise RuntimeError(f"{style_name} 未显式设置中文黑体")
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            if rfonts.get(qn(f"w:{attr}")) is not None:
                raise RuntimeError(f"{style_name} 仍包含主题字体属性 {attr}")

    print(
        f"generated: {path} | paragraphs={len(generated.paragraphs)} "
        f"tables={len(generated.tables)} headings={len(headings)}"
    )


validate_output(OUT)
